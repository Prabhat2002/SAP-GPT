import streamlit as st
from auth import send_otp_email, generate_otp
from database import (
    init_db, verify_otp, register_user_email, login_user_email, login_user_otp,
    email_exists, save_error, fetch_user_history, delete_error,
    create_team, fetch_user_teams, save_team_defect, send_join_request,
    fetch_join_requests, approve_request, reject_request,
    fetch_team_members, fetch_team_defects, search_similar_team_defects,
    delete_team, leave_team, save_feedback, count_pending_requests,
    get_username, fetch_user_feedback_stats
)
from backend import analyze_error, analyze_document
from report_generator import generate_report

import io
import os
import pandas as pd
import plotly.express as px
import PyPDF2

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from docx import Document

try:
    import pytesseract
    from PIL import Image, ImageEnhance
    if os.name == "nt":
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    TESSERACT_AVAILABLE = True
except Exception:
    TESSERACT_AVAILABLE = False

st.set_page_config(page_title="SAP-GPT", page_icon="🤖", layout="wide")

CSS = """
body { background: #020617; }
.main-title { font-size: 4rem; text-align: center; color: white; font-weight: 700; }
.subtitle { text-align: center; color: #cbd5e1; font-size: 1.2rem; }
.live-text { text-align: center; color: #06b6d4; font-size: 1rem; margin-bottom: 30px; }
.stButton > button {
    background: linear-gradient(90deg, #7c3aed, #06b6d4);
    color: white; border: none; border-radius: 12px;
    padding: 10px 20px; font-weight: 600;
}
"""
st.markdown(f"<style>{CSS}</style>", unsafe_allow_html=True)

init_db()

defaults = {
    "logged_in": False, "user_id": None,
    "auth_mode": "login",
    "otp_email": "",
    "otp_purpose": "",
    "pending_password": "",
    "chat_history": [],
    "last_analysis": None, "followups": [], "feedback_given": False,
    "similar_matches": [], "reg_success": False,
    "doc_analysis_result": None,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

st.markdown("""
<h1 class='main-title'>🤖 SAP-GPT</h1>
<p class='subtitle'>Enterprise AI Assistant for SAP Developers</p>
<p class='live-text'>⚡ ABAP · BTP · CAPM · Fiori · HANA · Basis ⚡</p>
""", unsafe_allow_html=True)

# =====================================================
# AUTH
# =====================================================
if not st.session_state.logged_in:

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:

        if st.session_state.auth_mode == "otp_verify":
            purpose = st.session_state.get("otp_purpose", "login")
            if purpose == "register":
                st.markdown("## 📧 Verify Your Email")
                st.info(f"OTP sent to **{st.session_state.otp_email}** — enter below to complete registration.")
            else:
                st.markdown("## 📧 Enter OTP")
                st.info(f"OTP sent to **{st.session_state.otp_email}** (valid 10 min)")

            otp_input = st.text_input("🔢 Enter 6-digit OTP", max_chars=6, placeholder="123456")

            if st.button("✅ Verify & Continue"):
                if verify_otp(st.session_state.otp_email, otp_input.strip()):
                    if purpose == "register":
                        ok = register_user_email(
                            st.session_state.otp_email,
                            st.session_state.get("pending_password", "")
                        )
                        if ok:
                            st.session_state.auth_mode = "login"
                            st.session_state.reg_success = True
                            st.session_state.otp_email = ""
                            st.session_state.pending_password = ""
                            st.rerun()
                        else:
                            st.error("❌ Email already registered. Please login.")
                    else:
                        user_id = login_user_otp(st.session_state.otp_email)
                        st.session_state.logged_in = True
                        st.session_state.user_id = user_id
                        st.session_state.auth_mode = "login"
                        st.success("✅ Login Successful!")
                        st.rerun()
                else:
                    st.error("❌ Invalid or expired OTP. Please try again.")

            if st.button("🔄 Resend OTP"):
                _, success, err = send_otp_email(st.session_state.otp_email)
                if err == "EMAIL_NOT_CONFIGURED":
                    st.warning("⚠️ Email not configured.")
                elif success:
                    st.success("✅ OTP resent!")
                else:
                    st.error(f"❌ {err}")

            if st.button("⬅ Back"):
                st.session_state.auth_mode = "register" if purpose == "register" else "login"
                st.rerun()

        elif st.session_state.auth_mode == "register":
            st.markdown("## ✨ Create Account")
            reg_email = st.text_input("📧 Email Address", key="reg_email")
            reg_pass  = st.text_input("🔒 Password", type="password", key="reg_pass")
            reg_conf  = st.text_input("🔒 Confirm Password", type="password", key="reg_conf")
            if reg_conf:
                if reg_pass != reg_conf:
                    st.error("❌ Passwords do not match")
                else:
                    st.success("✅ Passwords match")

            if st.button("📧 Register — Send OTP to Email"):
                if not reg_email.strip():
                    st.error("❌ Please enter your email")
                elif not reg_pass:
                    st.error("❌ Please enter a password")
                elif reg_pass != reg_conf:
                    st.error("❌ Passwords do not match")
                elif email_exists(reg_email.strip()):
                    st.error("❌ Email already registered. Please login.")
                else:
                    _, success, err = send_otp_email(reg_email.strip())
                    st.session_state.otp_email = reg_email.strip()
                    st.session_state.otp_purpose = "register"
                    st.session_state.pending_password = reg_pass
                    st.session_state.auth_mode = "otp_verify"
                    if err == "EMAIL_NOT_CONFIGURED":
                        st.warning("⚠️ Email service not configured.")
                    elif success:
                        st.success(f"✅ OTP sent to **{reg_email.strip()}**!")
                    st.rerun()

            st.markdown("---")
            if st.button("⬅ Back To Login"):
                st.session_state.auth_mode = "login"
                st.rerun()

        else:
            st.markdown("## 🚀 Welcome Back")
            if st.session_state.get("reg_success"):
                st.success("🎉 Registration successful! Please log in.")
                st.session_state.reg_success = False

            login_email = st.text_input("📧 Email Address", key="login_email")
            login_tab = st.radio("Sign in with:", ["🔒 Password", "🔢 OTP"], horizontal=True)

            if login_tab == "🔒 Password":
                login_pass = st.text_input("🔒 Password", type="password", key="login_pass")
                if st.button("🚀 Login"):
                    if not login_email.strip() or not login_pass:
                        st.error("❌ Please fill all fields")
                    else:
                        uid = login_user_email(login_email.strip(), login_pass)
                        if uid:
                            st.session_state.logged_in = True
                            st.session_state.user_id = uid
                            st.success("✅ Login Successful")
                            st.rerun()
                        else:
                            if email_exists(login_email.strip()):
                                st.error("❌ Wrong password.")
                            else:
                                st.error("❌ Email not registered. Please create an account first.")
            else:
                if st.button("📧 Send OTP"):
                    if not login_email.strip():
                        st.error("❌ Please enter your email")
                    elif not email_exists(login_email.strip()):
                        st.error("❌ Email not registered. Please create an account first.")
                    else:
                        _, success, err = send_otp_email(login_email.strip())
                        st.session_state.otp_email = login_email.strip()
                        st.session_state.otp_purpose = "login"
                        st.session_state.auth_mode = "otp_verify"
                        if err == "EMAIL_NOT_CONFIGURED":
                            st.warning("⚠️ Email not configured.")
                        elif success:
                            st.success("✅ OTP sent!")
                        st.rerun()

            st.markdown("---")
            st.markdown("### 🆕 New here?")
            if st.button("✨ Create Account"):
                st.session_state.auth_mode = "register"
                st.rerun()

# =====================================================
# MAIN APP
# =====================================================
else:
    _username = get_username(st.session_state.user_id)
    _history  = fetch_user_history(st.session_state.user_id)
    _teams    = fetch_user_teams(str(st.session_state.user_id))
    _fb_stats = fetch_user_feedback_stats(st.session_state.user_id)

    _spacer, _profile_col = st.columns([10, 1])
    with _profile_col:
        with st.popover(f"👤 {_username}"):
            st.markdown(f"## 👤 {_username}")
            st.markdown("---")
            _c1, _c2, _c3 = st.columns(3)
            _c1.metric("🔍 Analyses", len(_history))
            _c2.metric("🤝 Teams", len(_teams))
            _c3.metric("👍 Useful", _fb_stats.get("useful", 0))
            st.markdown("---")
            if _history:
                for _h in _history[:5]:
                    st.markdown(f"- 🖥 **{_h[2]}** — `{str(_h[3])[:40]}...`")
            else:
                st.caption("No saved analyses yet.")
            st.markdown("---")
            if st.button("🔓 Logout", key="profile_logout", type="primary"):
                for _k in list(st.session_state.keys()):
                    del st.session_state[_k]
                st.rerun()

    st.sidebar.markdown("### 🔔 Join Requests")
    if st.sidebar.button("🔄 Check Requests"):
        st.session_state["pending_count"] = count_pending_requests(str(st.session_state.user_id))
    _pending = st.session_state.get("pending_count", None)
    if _pending is None:
        st.sidebar.caption("Click above to check.")
    elif _pending > 0:
        st.sidebar.error(f"🔔 {_pending} pending request(s)!")
    else:
        st.sidebar.success("🟢 No pending requests")

    tabs = st.tabs(["💬 Chatbot", "📄 Document Analyzer", "🤝 Join Team", "🧠 Saved History", "📊 My Trends"])

    # =================================================
    # TAB 1 — CHATBOT
    # =================================================
    with tabs[0]:
        st.markdown("## 💬 SAP Chatbot")

        platform = st.selectbox("🖥 Select SAP Domain / Platform", [
            "SAP ABAP On-Premise",
            "SAP BTP (Business Technology Platform)",
            "SAP CAPM (Cloud Application Programming Model)",
            "SAP Fiori / UI5",
            "SAP HANA",
            "SAP Basis",
            "SAP Functional (MM / SD / FI / HR)",
            "SAP Integration Suite (CPI)",
            "SAP S/4HANA",
            "SAP RAP (RESTful ABAP Programming)",
            "Other / Not Sure",
        ], key="chatbot_platform")

        error_text = st.text_area(
            "📋 Paste SAP Error / Dump / Logs / Question",
            height=250,
            placeholder="Paste your SAP error, short dump, log, or ask any SAP question...",
            key="chatbot_input"
        )

        uploaded_image = st.file_uploader("📸 Upload Screenshot (optional)", type=["png", "jpg", "jpeg"], key="chatbot_img")
        if uploaded_image:
            if TESSERACT_AVAILABLE:
                try:
                    image = Image.open(uploaded_image).convert("L")
                    image = ImageEnhance.Contrast(image).enhance(2.0)
                    extracted_text = pytesseract.image_to_string(image, config='--psm 6')
                    st.success("✅ OCR Extraction Successful")
                    ocr_edited = st.text_area("📝 OCR Extracted Text (edit if needed)", value=extracted_text, height=150, key="ocr_edit")
                    st.caption("💡 Tip: If OCR missed some text, manually paste it above.")
                    if not error_text:
                        error_text = ocr_edited
                    else:
                        error_text = error_text + "\n\n[From Screenshot]:\n" + ocr_edited
                except Exception as e:
                    st.warning(f"⚠️ OCR failed: {e}")
            else:
                st.info("ℹ️ Tesseract OCR not installed.")

        if st.button("⚡ Ask SAP-GPT"):
            if not error_text.strip():
                st.warning("⚠️ Please enter a question or error.")
            else:
                _all_teams = fetch_user_teams(str(st.session_state.user_id))
                _pre_similar = []
                for _pt in _all_teams:
                    _pm = search_similar_team_defects(_pt[1], error_text)
                    for _m in _pm:
                        _pre_similar.append((_pt[0], _pt[1], _m[0], _m[1], _m[2]))
                st.session_state.similar_matches = _pre_similar

                with st.spinner("🤖 SAP-GPT is thinking..."):
                    response = analyze_error(platform, error_text)

                st.session_state.last_analysis = (platform, error_text, response)
                st.session_state.followups = []
                st.session_state.feedback_given = False

        if st.session_state.get("similar_matches"):
            st.markdown("---")
            st.warning("⚠️ **This issue was already addressed in your team!**")
            for _tname, _tcode, _mid, _merr, _mans in st.session_state.similar_matches[:3]:
                with st.expander(f"🔁 Already solved in **{_tname}**", expanded=True):
                    st.markdown("**Original Question / Error:**")
                    st.code(_merr[:600])
                    st.markdown("**Team Solution / Analysis:**")
                    st.success(_mans[:1500])
            st.markdown("---")

        if st.session_state.get("last_analysis"):
            _p, _e, _r = st.session_state.last_analysis
            st.markdown("## 🎯 AI Response")
            st.markdown(_r)

        if st.session_state.last_analysis:
            platform_a, error_text_a, response_a = st.session_state.last_analysis

            if st.session_state.get("followups"):
                st.markdown("---")
                st.markdown("### 💬 Conversation")
                for fq, fa in st.session_state.followups:
                    with st.chat_message("user"):
                        st.markdown(fq)
                    with st.chat_message("assistant"):
                        st.markdown(fa)

            st.markdown("---")
            follow_up = st.chat_input("Ask a follow-up question...", key="followup_chat")
            if follow_up and follow_up.strip():
                context = (
                    f"Platform: {platform_a}\n\nOriginal Error:\n{error_text_a}\n\n"
                    f"Initial Analysis:\n{response_a}"
                )
                for fq, fa in st.session_state.get("followups", []):
                    context += f"\n\nUser Follow-up: {fq}\nAssistant: {fa}"
                context += f"\n\nNew Follow-up Question: {follow_up}"
                with st.spinner("🤖 Answering..."):
                    fu_response = analyze_error(platform_a, context, use_feedback=False)
                st.session_state.followups.append((follow_up, fu_response))
                st.rerun()

            severity = "MEDIUM"
            if "HIGH" in response_a: severity = "HIGH"
            if "CRITICAL" in response_a: severity = "CRITICAL"

            # ---- DOWNLOAD SECTION ----
            st.markdown("---")
            st.markdown("### 📥 Download Report")
            download_format = st.radio("Format", ["TXT", "PDF", "WORD"], horizontal=True)

            if download_format == "TXT":
                txt_content = f"""SAP-GPT ANALYSIS REPORT
{'='*60}
PLATFORM: {platform_a}
{'='*60}
ERROR / QUESTION:
{'='*60}
{error_text_a}

{'='*60}
FULL ANALYSIS:
{'='*60}
{response_a}
"""
                st.download_button("⬇ Download (TXT)", txt_content,
                                   file_name="sap_gpt_report.txt", mime="text/plain")

            elif download_format == "PDF":
                try:
                    pdf_buffer = io.BytesIO()
                    doc_pdf = SimpleDocTemplate(
                        pdf_buffer, pagesize=A4,
                        rightMargin=20*mm, leftMargin=20*mm,
                        topMargin=20*mm, bottomMargin=20*mm
                    )
                    styles = getSampleStyleSheet()
                    title_style = ParagraphStyle('CTitle', parent=styles['Title'],
                        fontSize=16, spaceAfter=12, textColor=colors.HexColor('#7c3aed'))
                    heading_style = ParagraphStyle('CHeading', parent=styles['Heading2'],
                        fontSize=12, spaceAfter=6, textColor=colors.HexColor('#1e293b'))
                    body_style = ParagraphStyle('CBody', parent=styles['Normal'],
                        fontSize=9, spaceAfter=4, leading=14)

                    story = []
                    story.append(Paragraph("SAP-GPT Analysis Report", title_style))
                    story.append(Spacer(1, 6*mm))
                    story.append(Paragraph(f"Platform: {platform_a}", heading_style))
                    story.append(Spacer(1, 4*mm))
                    story.append(Paragraph("Error / Question:", heading_style))
                    for line in error_text_a.split("\n"):
                        clean = line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                        if clean.strip():
                            story.append(Paragraph(clean, body_style))
                        else:
                            story.append(Spacer(1, 2*mm))
                    story.append(Spacer(1, 6*mm))
                    story.append(Paragraph("Full Analysis:", heading_style))
                    for line in response_a.split("\n"):
                        clean = line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                        if clean.strip():
                            if clean.startswith("##"):
                                story.append(Paragraph(f"<b>{clean.replace('#','').strip()}</b>", body_style))
                            else:
                                story.append(Paragraph(clean, body_style))
                        else:
                            story.append(Spacer(1, 2*mm))
                    doc_pdf.build(story)
                    pdf_buffer.seek(0)
                    st.download_button("⬇ Download (PDF)", pdf_buffer,
                                       file_name="sap_gpt_report.pdf", mime="application/pdf")
                except Exception as e:
                    st.error(f"❌ PDF generation failed: {e}")

            elif download_format == "WORD":
                try:
                    word_doc = Document()
                    word_doc.add_heading("SAP-GPT Analysis Report", 0)
                    word_doc.add_heading("Platform", level=2)
                    word_doc.add_paragraph(platform_a)
                    word_doc.add_heading("Error / Question", level=2)
                    word_doc.add_paragraph(error_text_a)
                    word_doc.add_heading("Full Analysis", level=2)
                    for line in response_a.split("\n"):
                        if line.startswith("## "):
                            word_doc.add_heading(line.replace("## ", ""), level=3)
                        elif line.startswith("# "):
                            word_doc.add_heading(line.replace("# ", ""), level=2)
                        elif line.strip():
                            word_doc.add_paragraph(line)
                    word_buffer = io.BytesIO()
                    word_doc.save(word_buffer)
                    word_buffer.seek(0)
                    st.download_button("⬇ Download (Word)", word_buffer,
                                       file_name="sap_gpt_report.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e:
                    st.error(f"❌ Word generation failed: {e}")

            st.markdown("---")
            if st.button("💾 Save For Later"):
                save_error(st.session_state.user_id, platform_a, error_text_a, response_a, severity)
                st.success("✅ Saved to history!")

            st.markdown("---")
            st.markdown("### 👥 Save To Team")
            teams = fetch_user_teams(str(st.session_state.user_id))
            if teams:
                team_dropdown = [f"{t[0]} ({t[1]})" for t in teams]
                selected_team = st.selectbox("📁 Select Team", team_dropdown)
                if st.button("🚀 Push To Team"):
                    code = selected_team.split("(")[-1].replace(")", "").strip()
                    save_team_defect(code, selected_team, error_text_a, response_a, str(st.session_state.user_id))
                    st.success("✅ Added to team!")
            else:
                st.info("ℹ️ No teams yet. Create one in Join Team tab.")

            st.markdown("---")
            st.markdown("### 💡 Was this helpful?")
            if not st.session_state.get("feedback_given"):
                col_good, col_bad, col_ok = st.columns(3)
                with col_good:
                    if st.button("👍 Useful", key="fb_useful"):
                        save_feedback(st.session_state.user_id, platform_a, error_text_a, response_a, "useful")
                        st.session_state.feedback_given = True
                        st.rerun()
                with col_bad:
                    if st.button("👎 Not Helpful", key="fb_bad"):
                        save_feedback(st.session_state.user_id, platform_a, error_text_a, response_a, "not_useful")
                        st.session_state.feedback_given = True
                        st.rerun()
                with col_ok:
                    if st.button("😐 Partial", key="fb_ok"):
                        save_feedback(st.session_state.user_id, platform_a, error_text_a, response_a, "partial")
                        st.session_state.feedback_given = True
                        st.rerun()
            else:
                st.success("✅ Feedback submitted.")

    # =================================================
    # TAB 2 — DOCUMENT ANALYZER
    # =================================================
    with tabs[1]:
        st.markdown("## 📄 SAP Document Analyzer")
        st.caption("Upload any SAP document — TSD, FSD, Defect, or any SAP-related file.")

        uploaded_doc = st.file_uploader("📂 Upload Document (PDF, Word, TXT)", type=["pdf", "docx", "txt"], key="doc_uploader")

        doc_type = st.selectbox("📋 Document Type", [
            "TSD (Technical Specification)",
            "FSD (Functional Specification)",
            "Defect Report",
            "Other SAP Related Document",
        ], key="doc_type_select")

        if doc_type == "Other SAP Related Document":
            st.info("ℹ️ Upload any SAP doc: ABAP, RAP, BTP, Fiori, HANA, CPI, Functional, Basis, etc. Non-SAP documents will be rejected.")

        if st.button("⚡ Analyze Document"):
            if not uploaded_doc:
                st.warning("⚠️ Please upload a document first.")
            else:
                extracted_text = ""
                with st.spinner("📖 Reading document..."):
                    try:
                        if uploaded_doc.name.endswith(".txt"):
                            extracted_text = uploaded_doc.read().decode("utf-8", errors="ignore")
                        elif uploaded_doc.name.endswith(".docx"):
                            d = Document(uploaded_doc)
                            extracted_text = "\n".join(p.text for p in d.paragraphs)
                        elif uploaded_doc.name.endswith(".pdf"):
                            reader = PyPDF2.PdfReader(uploaded_doc)
                            for page in reader.pages:
                                txt = page.extract_text()
                                if txt:
                                    extracted_text += txt
                    except Exception as e:
                        st.error(f"❌ Could not read document: {e}")

                if extracted_text.strip():
                    with st.spinner("🤖 Analyzing document..."):
                        doc_response = analyze_document(doc_type, extracted_text)
                    st.session_state.doc_analysis_result = (doc_type, uploaded_doc.name, doc_response)
                    st.rerun()
                else:
                    st.warning("⚠️ Could not extract text. Try a different file format.")

        if st.session_state.get("doc_analysis_result"):
            dtype, dname, dresult = st.session_state.doc_analysis_result
            if "does not appear to be SAP-related" in dresult:
                st.error(dresult)
            else:
                st.markdown("---")
                st.markdown("## 📊 Document Analysis Result")
                st.caption(f"📄 **{dname}** — {dtype}")
                st.markdown(dresult)

    # =================================================
    # TAB 3 — JOIN TEAM
    # =================================================
    with tabs[2]:
        st.markdown("## 🤝 Team Collaboration")

        with st.expander("🚀 Create a New Team", expanded=True):
            team_name    = st.text_input("👥 Team Name")
            project_name = st.text_input("📁 Project Name")
            notes        = st.text_area("📝 Notes (optional)")
            if st.button("✅ Create Team"):
                if not team_name.strip():
                    st.warning("⚠️ Please enter a team name.")
                else:
                    team_code = create_team(team_name, project_name, str(st.session_state.user_id), notes)
                    st.success(f"✅ Team **{team_name}** created!")
                    st.info(f"🔑 Team Code: **`{team_code}`**")

        st.markdown("---")
        st.markdown("## 📋 My Teams")
        teams = fetch_user_teams(str(st.session_state.user_id))
        if teams:
            for t in teams:
                t_name, t_code, t_owner = t[0], t[1], t[2]
                is_owner = str(t_owner) == str(st.session_state.user_id)
                role_tag = "👑 Owner" if is_owner else "👤 Member"
                with st.expander(f"🚀 {t_name}  |  {role_tag}  |  Code: {t_code}", expanded=True):
                    if is_owner:
                        st.info(f"🔑 Share code: **`{t_code}`**")
                    else:
                        st.success(f"✅ Member (Code: `{t_code}`)")

                    st.markdown("### 👥 Members")
                    members = fetch_team_members(t_code)
                    for m in members:
                        icon = "👑" if m[2] == "owner" else "👤"
                        st.markdown(f"- {icon} **{m[1]}**")

                    st.markdown("### 🐛 Shared Issues")
                    defects = fetch_team_defects(t_code)
                    if defects:
                        for d in defects:
                            with st.expander(f"📌 Issue #{d[0]}"):
                                st.code(d[2])
                                st.write(d[3])
                    else:
                        st.caption("No issues pushed yet.")

                    st.markdown("---")
                    if is_owner:
                        confirm_key = f"confirm_del_{t_code}"
                        if st.session_state.get(confirm_key):
                            c1, c2 = st.columns(2)
                            with c1:
                                if st.button("🔴 Yes, Delete", key=f"yes_del_{t_code}"):
                                    delete_team(t_code, str(st.session_state.user_id))
                                    st.session_state[confirm_key] = False
                                    st.rerun()
                            with c2:
                                if st.button("↩ Cancel", key=f"no_del_{t_code}"):
                                    st.session_state[confirm_key] = False
                                    st.rerun()
                        else:
                            if st.button(f"🗑 Delete Team", key=f"del_team_{t_code}"):
                                st.session_state[confirm_key] = True
                                st.rerun()
                    else:
                        leave_key = f"confirm_leave_{t_code}"
                        if st.session_state.get(leave_key):
                            c1, c2 = st.columns(2)
                            with c1:
                                if st.button("🔴 Yes, Leave", key=f"yes_leave_{t_code}"):
                                    leave_team(t_code, str(st.session_state.user_id))
                                    st.session_state[leave_key] = False
                                    st.rerun()
                            with c2:
                                if st.button("↩ Cancel", key=f"no_leave_{t_code}"):
                                    st.session_state[leave_key] = False
                                    st.rerun()
                        else:
                            if st.button(f"🚪 Leave Team", key=f"leave_{t_code}"):
                                st.session_state[leave_key] = True
                                st.rerun()
        else:
            st.info("ℹ️ No teams yet.")

        st.markdown("---")
        st.markdown("## 🔗 Join Existing Team")
        join_code = st.text_input("🔑 Enter Team Code")
        if st.button("📨 Send Join Request"):
            if not join_code.strip():
                st.warning("⚠️ Please enter a code.")
            else:
                send_join_request(join_code.strip().upper(), str(st.session_state.user_id))
                st.success("✅ Request sent!")

        st.markdown("---")
        st.markdown("## 🛡 Pending Join Requests")
        requests = fetch_join_requests(str(st.session_state.user_id))
        if requests:
            for req in requests:
                req_id, team_code_r, user_id_r, display_name, team_name_r = req
                st.markdown(f"👤 **{display_name}** → **{team_name_r}** (`{team_code_r}`)")
                c1, c2, _ = st.columns([1, 1, 4])
                with c1:
                    if st.button("✅ Approve", key=f"approve_{req_id}"):
                        approve_request(req_id, team_code_r, user_id_r)
                        st.rerun()
                with c2:
                    if st.button("❌ Reject", key=f"reject_{req_id}"):
                        reject_request(req_id)
                        st.rerun()
                st.divider()
        else:
            st.info("ℹ️ No pending requests.")

    # =================================================
    # TAB 4 — SAVED HISTORY
    # =================================================
    with tabs[3]:
        st.markdown("## 🧠 Saved History")
        rows = fetch_user_history(st.session_state.user_id)
        if rows:
            for row in rows:
                with st.expander(f"🖥 {row[2]}  |  #{row[0]}"):
                    st.code(row[3])
                    st.write(row[4])
                    if st.button(f"🗑 Delete", key=f"del_{row[0]}"):
                        delete_error(row[0])
                        st.rerun()
        else:
            st.info("ℹ️ Nothing saved yet.")

    # =================================================
    # TAB 5 — TRENDS
    # =================================================
    with tabs[4]:
        st.markdown("## 📊 My SAP Trends")
        rows = fetch_user_history(st.session_state.user_id)
        if rows:
            df = pd.DataFrame([{"platform": r[2], "severity": r[5]} for r in rows])
            col1, col2 = st.columns(2)
            with col1:
                fig = px.pie(df, names="severity", title="Severity Distribution",
                             color_discrete_sequence=px.colors.sequential.Plasma)
                st.plotly_chart(fig, use_container_width=True)
            with col2:
                fig2 = px.histogram(df, x="platform", title="Errors By Platform",
                                    color_discrete_sequence=["#7c3aed"])
                st.plotly_chart(fig2, use_container_width=True)
        else:
            st.info("ℹ️ No data yet.")